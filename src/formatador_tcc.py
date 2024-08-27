from docx import Document
from docx.shared import Pt, Inches

def formatar_tcc(texto):
    # Criar um novo documento
    doc = Document()
    
    # Formatar título
    titulo = doc.add_heading('Título do TCC', level=1)
    titulo.alignment = 1  # Centralizado
    
    # Adicionar texto formatado
    paragrafo = doc.add_paragraph(texto)
    paragrafo.style.font.name = 'Arial'
    paragrafo.style.font.size = Pt(12)
    paragrafo.paragraph_format.line_spacing = Pt(18)  # 1,5 espaços
    
    # Configurar margens
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.18)  # 3 cm
        section.bottom_margin = Inches(0.79)  # 2 cm
        section.left_margin = Inches(1.18)  # 3 cm
        section.right_margin = Inches(0.79)  # 2 cm

    # Salvar documento
    doc.save('TCC_formatado.docx')

# Uso da função
if __name__ == "__main__":
    texto_tcc = "Seu texto aqui..."
    formatar_tcc(texto_tcc)
